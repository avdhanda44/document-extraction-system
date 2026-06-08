from docx import Document


def extract_text_from_docx(docx_path):
    # DOCX files can have text in paragraphs and tables.
    # We read both so form values are not missed.
    document = Document(docx_path)
    text_parts = []

    # Read normal paragraphs.
    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.strip()

        if paragraph_text:
            text_parts.append(paragraph_text)

    # Read table cells too.
    for table in document.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]

            if row_text:
                text_parts.append(" ".join(row_text))

    final_text = "\n".join(text_parts).strip()

    if final_text == "":
        raise ValueError("No readable text found in DOCX file.")

    return final_text
