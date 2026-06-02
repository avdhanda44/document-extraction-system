# --------------------------------------------------------------------------------
# # Text Extraction
# 
# This notebook figures out what kind of file was uploaded and turns that file into plain text for the field extractor.
# --------------------------------------------------------------------------------

import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory

import easyocr
import pdfplumber
from docx import Document
from pdf2image import convert_from_path

# Detect file type
# We check the file content, not just the file extension, because extensions can be renamed by mistake.
def detect_file_type(file_path):
    with open(file_path, "rb") as file:
        file_header = file.read(10)

    if file_header.startswith(b"%PDF"):
        return "pdf"

    if file_header.startswith(b"\x89PNG"):
        return "png"

    if file_header.startswith(b"\xff\xd8"):
        return "jpg"

    if zipfile.is_zipfile(file_path):
        with zipfile.ZipFile(file_path) as docx_file:
            if "word/document.xml" in docx_file.namelist():
                return "docx"

    return "unknown"

# Extract raw text
# Digital PDFs are read directly with pdfplumber.
# Scanned PDFs and image files go through EasyOCR.
# DOCX files are read with python-docx.
reader = None


def get_ocr_reader():
    global reader

    if reader is None:
        reader = easyocr.Reader(['en'], gpu=False)

    return reader


def ocr_image_file(image_path):
    ocr_reader = get_ocr_reader()
    results = ocr_reader.readtext(str(image_path))

    return "\n".join(item[1] for item in results)


def extract_scanned_pdf_text(file_path):
    try:
        pages = convert_from_path(str(file_path), dpi=350)
    except Exception as error:
        raise RuntimeError(
            "Scanned PDF OCR needs Poppler installed for pdf2image."
        ) from error

    page_text = []

    with TemporaryDirectory() as temp_folder:
        temp_folder = Path(temp_folder)

        for page_number, page in enumerate(pages, start=1):
            page_image_path = temp_folder / f"page_{page_number}.png"
            page.save(page_image_path, "PNG")
            page_text.append(ocr_image_file(page_image_path))

    final_text = "\n".join(page_text).strip()

    if final_text == "":
        raise ValueError("No readable text found after scanned PDF OCR.")

    return final_text


def extract_docx_text(file_path):
    document = Document(file_path)
    text_parts = []

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.strip()
        if paragraph_text:
            text_parts.append(paragraph_text)

    for table in document.tables:
        for row in table.rows:
            row_values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_values:
                text_parts.append(" ".join(row_values))

    final_text = "\n".join(text_parts).strip()

    if final_text == "":
        raise ValueError("No readable text found in DOCX file.")

    return final_text


def extract_text(file_path, file_type):
    if file_type == "pdf":
        with pdfplumber.open(file_path) as pdf:
            page_text = [page.extract_text() or "" for page in pdf.pages]

        final_text = "\n".join(page_text).strip()

        if final_text:
            return final_text

        return extract_scanned_pdf_text(file_path)

    if file_type in ["png", "jpg"]:
        return ocr_image_file(file_path)

    if file_type == "docx":
        return extract_docx_text(file_path)

    raise ValueError(f"Unsupported file type: {file_type}")

